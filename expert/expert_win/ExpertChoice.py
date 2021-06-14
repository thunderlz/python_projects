# coding:utf-8
# coding by leiz

import tkinter as tk
# from tkinter import *
from tkinter import Button, Label, Entry, Tk, PhotoImage, Frame, StringVar, W, S, EW, VERTICAL, NSEW, NS, Listbox, \
    MULTIPLE, Toplevel, END, E
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime as dt
from datetime import timedelta as delta
from sys import exit

from docx import Document
import xlrd
from xlutils.copy import copy
import xlwt

import sqlite3
import numpy as np
# from numpy.core import _dtype_ctypes
import pandas as pd
from pandas import DataFrame, read_sql, merge, concat, ExcelWriter, ExcelFile
from numpy import nan


# from PIL import Image, ImageTk
class expertChoice:
    def __init__(self):
        # 统一的对齐方式
        self.duiqi = 'W'
        # 控件的宽
        self.cbwd = 27
        self.etwd = 29
        self.letwd = 73
        self.stwd = 10
        self.root = Tk()
        self.root.title("中捷专家抽取系统")
        self.root.resizable(0, 0)
        self.root.geometry('960x640+200+100')
        self.root.iconbitmap('logo.ico')

        # 全局变量
        # 由于要考虑多次抽取的问题，所以定义一个抽取结果的全局变量
        self.dfrltexpert = DataFrame()
        # 抽取的批次
        self.choicetime_int = 0
        self.conditiontre_dict = {}

        self.fieldlist = ['专业类别', '工作单位', '工作部门', '区域', '专业资质等级']

        # 业主代表的信息
        self.bossname = ''
        self.bossworkpart = ''
        self.bosstelno = ''
        self.bossmobilno = ''
        self.bossemail = ''
        self.bossnaid = ''
        self.dfboss = pd.DataFrame([], index=range(0, 10),
                                   columns=['bossname', 'bossworkpart', 'bosstelno', 'bossmobilno', 'bossemail',
                                            'bossnaid'])
        self.dfboss[:] = '未填写'

        # 初始化数据库
        self.dbconn = sqlite3.connect('mydb.db')
        self.dbcur = self.dbconn.cursor()
        # self.dbcur.execute('create table if not exists tbexpert(id)')
        # self.dbcur.execute('create table if not exists tbcata(id)')

        # 表头
        self.topimg = PhotoImage(file='top.gif')
        self.banner = ttk.Label(self.root, image=self.topimg)
        self.banner.grid(row=0, column=0, sticky='NWSE')

        # 四个框架
        self.base = Frame(self.root, width=960, height=200)
        self.show = Frame(self.root, width=960, height=120)
        self.rlt = Frame(self.root, width=960, height=160)
        self.manage = Frame(self.root, width=960, height=100)

        # 定义基本信息区域
        # 第一行
        Label(self.base, text='招标组织方式:').grid(row=0, column=0, sticky=self.duiqi)
        self.org_var = StringVar()
        self.org_cb = ttk.Combobox(self.base, textvariable=self.org_var, width=self.cbwd)
        self.org_cb['values'] = ['公开招标', '邀请招标', '公开询价', '单一来源采购']
        self.org_cb.current(0)
        self.org_cb.grid(row=0, column=1, sticky=W)

        Label(self.base, text='招标人名称:').grid(row=0, column=2, sticky=self.duiqi)
        self.bossname_var = StringVar()
        self.boss = Entry(self.base, width=self.etwd, textvariable=self.bossname_var)
        self.boss.grid(row=0, column=3, sticky=self.duiqi)

        Label(self.base, text='招标代理机构\n项目编号:', width=15).grid(row=0, column=4, sticky=self.duiqi)
        self.agentid = Entry(self.base, width=self.etwd)
        self.agentid.grid(row=0, column=5, sticky=W)

        # 第二行
        Label(self.base, text='项目编号:').grid(row=1, column=0, sticky=self.duiqi)
        self.projectid_var = StringVar()
        self.projectid = Entry(self.base, width=self.letwd, textvariable=self.projectid_var)
        self.projectid.grid(row=1, column=1, columnspan=3, sticky=self.duiqi)

        Label(self.base, text='招标代理机构\n项目名称:', width=15).grid(row=1, column=4, sticky=self.duiqi)
        self.agentnames = StringVar()
        self.agentname_cb = ttk.Combobox(self.base, height=1, textvariable=self.agentnames, width=self.cbwd)
        self.agentname_cb['values'] = ('中捷通信有限公司', '公诚咨询管理有限公司', '不委托')
        self.agentname_cb.current(0)
        self.agentname_cb.grid(row=1, column=5, sticky=self.duiqi)

        # 第三行
        Label(self.base, text='项目名称:').grid(row=2, column=0, sticky=self.duiqi)
        self.projectname_var = StringVar()
        self.projectname = Entry(self.base, width=self.letwd, textvariable=self.projectname_var)
        self.projectname.grid(row=2, column=1, columnspan=3, sticky=self.duiqi)

        Label(self.base, text='递补专家人数:', width=15).grid(row=2, column=4, sticky=self.duiqi)
        self.backexpertnum_var = StringVar()
        self.backexpertnum_cb = ttk.Combobox(self.base, height=1, state='readonly', textvariable=self.backexpertnum_var,
                                             width=self.cbwd)
        self.backexpertnum_cb['values'] = (0, 1, 2, 3)
        self.backexpertnum_cb.current(0)
        self.backexpertnum_cb.grid(row=2, column=5, sticky=self.duiqi)

        # 第四行
        Label(self.base, text='抽取地点:').grid(row=3, column=0, sticky=self.duiqi)
        self.choiceplace_var = StringVar()
        self.choiceplace = Entry(self.base, width=self.letwd, textvariable=self.choiceplace_var)
        self.choiceplace.grid(row=3, column=1, columnspan=3, sticky=self.duiqi)

        # 抽取时间
        Label(self.base, text='抽取时间:', width=15).grid(row=3, column=4, sticky=self.duiqi)
        self.choicedate_var = StringVar()
        self.choicedate_cb = ttk.Combobox(self.base, height=1, state='readonly', textvariable=self.choicedate_var,
                                          width=self.cbwd)
        self.choicedate_cb['values'] = [(delta(day) + dt.now()).strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')
                                        for day in range(-10, 10)]
        self.choicedate_cb.current(10)
        self.choicedate_cb.grid(row=3, column=5, sticky=self.duiqi)

        # 第五行
        Label(self.base, text='评标委员会人数:').grid(row=4, column=0, sticky=self.duiqi)
        self.confnum_var = StringVar()
        self.confnum_cb = ttk.Combobox(self.base, textvariable=self.confnum_var, state='readonly', width=self.cbwd)
        self.confnum_cb['values'] = (5, 7, 9, 11, 13, 15, 17, 19)
        self.confnum_cb.current(0)
        self.confnum_cb.grid(row=4, column=1, sticky=self.duiqi)
        # self.confnum_et.bind('<FocusOut>', self.checkconfnum())

        Label(self.base, text='业主代表人数:').grid(row=4, column=2, sticky=self.duiqi)
        self.bossnum_var = StringVar()
        self.bossnum_cb = ttk.Combobox(self.base, state='readonly', height=1, textvariable=self.bossnum_var,
                                       width=self.cbwd)
        self.bossnum_cb['values'] = (0, 1, 2, 3, 4)
        self.bossnum_cb.current(1)
        self.bossnum_cb.grid(row=4, column=3, sticky=self.duiqi)

        # 评审时间
        Label(self.base, text='评审时间:', width=15).grid(row=4, column=4, sticky=self.duiqi)
        self.meetdate_var = StringVar()
        self.meetdate_cb = ttk.Combobox(self.base, height=1, textvariable=self.meetdate_var, width=self.cbwd)
        self.meetdate_cb['values'] = [(delta(day) + dt.now()).strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')
                                      for day in range(-10, 10)]
        self.meetdate_cb.current(10)
        self.meetdate_cb.grid(row=4, column=5, sticky=self.duiqi)

        # 第六行
        Label(self.base, text='专家评委人数:').grid(row=5, column=0, sticky=self.duiqi)
        self.expertnum_var = StringVar()
        self.expertnum = Entry(self.base, state='readonly', bg='#808080', textvariable=self.expertnum_var,
                               width=self.etwd)
        self.expertnum.grid(row=5, column=1, sticky=self.duiqi)

        Label(self.base, text='需抽取专家\n（含递补）人数:').grid(row=5, column=2, sticky=self.duiqi)
        self.expertchoicenum_var = StringVar()
        self.expertchoicenum = Entry(self.base, state='readonly', textvariable=self.expertchoicenum_var,
                                     width=self.etwd)
        self.expertchoicenum.grid(row=5, column=3, sticky=self.duiqi)

        # 业主代表的信息
        self.bossname_lb = Label(self.base)
        self.bossname_lb.grid(row=5, column=5, sticky=S)

        self.bossinfo_btn = Button(self.base, text='填写业主\n代表信息', command=self.bossinfo_func)
        self.bossinfo_btn.grid(row=5, column=4, columnspan=1, rowspan=1, ipadx=0, ipady=0, sticky=EW)

        # 展示抽取条件列表区域
        self.tree = ttk.Treeview(self.show, show="headings", height=5, columns=("a", "b"))
        self.vbar = ttk.Scrollbar(self.show, orient=VERTICAL, command=self.tree.yview)
        # 定义树形结构与滚动条
        self.tree.configure(yscrollcommand=self.vbar.set)
        # 表格的标题
        self.tree.column("a", width=15 * self.stwd, anchor="center")
        self.tree.column("b", width=16 * self.stwd, anchor="center")
        self.tree.heading("a", text="筛选字段")
        self.tree.heading("b", text="筛选条件")
        # 调用方法获取表格内容插入
        self.tree.grid(row=0, column=0, sticky=NSEW, rowspan=4)
        self.vbar.grid(row=0, column=1, sticky=NS, rowspan=4)

        # 展示筛选条件
        Label(self.show, text='字段名称:').grid(row=0, column=2, padx=50, sticky=self.duiqi)
        self.choicefield_var = StringVar()
        self.choicefield_cb = ttk.Combobox(self.show, width=10, height=1, state='readonly',
                                           textvariable=self.choicefield_var)
        self.choicefield_cb.grid(row=0, column=3, sticky=self.duiqi, padx=int(0.5 * self.stwd))
        self.choicefield_cb['values'] = self.fieldlist
        self.choicefield_cb.current(0)

        self.condtion_lb = Label(self.show, text='筛选条件:')
        self.condtion_lb.grid(row=1, column=2)
        self.choicecondition1_var = StringVar()
        self.choicecondition1_cb = ttk.Combobox(self.show, width=10, height=1, state='readonly',
                                                textvariable=self.choicecondition1_var)
        self.choicecondition1_cb.grid(row=1, column=3, sticky=self.duiqi, padx=int(0.5 * self.stwd))
        # self.choicecondition1_cb.grid_forget()

        # Label(self.show, text='筛选条件2').grid(row=2, column=2)
        # self.choicecondition2_var = StringVar()
        # self.choicecondition2_cb = ttk.Combobox(self.show, height=1, textvariable=self.choicecondition2_var)
        # # self.choicecondition2_cb['values'] = (1, 2, 3)
        # self.choicecondition2_cb.grid(row=2, column=3,sticky=self.duiqi)

        self.addcondition_btn = Button(self.show, text='添加筛选条件', height=1, width=int(self.stwd * 1.2),
                                       command=self.addcondition_func)
        self.addcondition_btn.grid(row=3, column=3)
        self.delcondition_btn = Button(self.show, text='清空筛选条件', height=1, width=int(self.stwd * 1.2),
                                       command=self.delcondition_func)
        self.delcondition_btn.grid(row=3, column=2)

        self.choicecondition_lbvar = StringVar()
        self.choicecondition_lb = Listbox(self.show, height=6, selectmode=MULTIPLE,
                                          listvariable=self.choicecondition_lbvar)
        self.choicecondition_lb.grid(row=0, column=4, rowspan=4, padx=10, sticky=self.duiqi)

        Button(self.show, text='抽取专家', height=5, width=int(self.stwd * 1.5), font=20,
               command=self.choiceexpert_func).grid(row=0, column=5, rowspan=4, padx=10, columnspan=1)

        # 抽取结果
        self.rlt_lb = Label(self.rlt, text='目前已抽取人数：0', width=68, anchor='w')
        self.rlt_lb.grid(row=0, column=0, sticky=self.duiqi)

        # 导出结果按钮
        self.export = Button(self.rlt, text='导出结果', width=15, command=self.exportrlt_func)
        self.export.grid(row=0, column=2, sticky='E', columnspan=2)
        # 重置结果按钮
        self.reset_btn = Button(self.rlt, text='重置结果', width=15, command=self.resetrlttree)
        self.reset_btn.grid(row=0, column=1, sticky='E')

        self.rlttree = ttk.Treeview(self.rlt, show="headings", height=5,
                                    columns=("i", "a", "b", "c", "d", "e", "f", 'g', 'h'))
        self.rltvbar = ttk.Scrollbar(self.rlt, orient=VERTICAL, command=self.rlttree.yview)
        # 定义树形结构与滚动条
        self.rlttree.configure(yscrollcommand=self.rltvbar.set)
        # 表格的标题
        self.rlttree.column("i", width=50, anchor="center")
        self.rlttree.column("a", width=80, anchor="center")
        self.rlttree.column("b", width=50, anchor="center")
        self.rlttree.column("c", width=180, anchor="center")
        self.rlttree.column("d", width=190, anchor="center")
        self.rlttree.column("e", width=50, anchor="center")
        self.rlttree.column("f", width=130, anchor="center")
        self.rlttree.column("g", width=90, anchor="center")
        self.rlttree.column("h", width=90, anchor="center")
        self.rlttree.heading("i", text="序号")
        self.rlttree.heading("a", text="姓名")
        self.rlttree.heading("b", text="性别")
        self.rlttree.heading("c", text='工作单位')
        self.rlttree.heading("d", text="工作部门")
        self.rlttree.heading("e", text="区域")
        self.rlttree.heading("f", text="资质等级")
        self.rlttree.heading("g", text="是否参加")
        self.rlttree.heading("h", text="抽取批次")
        # 调用方法获取表格内容插入
        self.rlttree.grid(row=1, column=0, sticky='NEW', columnspan=3)
        self.rltvbar.grid(row=1, column=3, sticky='NS', columnspan=3)

        # 管理区域
        self.importdata_btn = Button(self.manage, text='导入专家库', command=self.importdata_func)
        self.importdata_btn.grid(row=0, column=0, ipadx=30, padx=0, pady=5, sticky='W')

        self.cleardata_btn = Button(self.manage, text='清空专家库', command=self.cleardata_func)
        self.cleardata_btn.grid(row=0, column=1, ipadx=30, padx=0, sticky='W')

        self.checkdata_btn = Button(self.manage, text='查看专家库', command=self.showexpert_func)
        self.checkdata_btn.grid(row=0, column=2, ipadx=30, padx=0, sticky='W')

        self.checkdata_btn = Button(self.manage, text='查看历史', command=self.showhistory_func)
        self.checkdata_btn.grid(row=0, column=3, ipadx=30, padx=0, sticky='W')

        self.exit_btn = Button(self.manage, text='退出程序', width=30, command=exit)
        self.exit_btn.grid(row=0, column=3, ipadx=30, padx=250, sticky='E')

        # 整体区域定位
        self.base.grid(row=1, column=0, pady=3, padx=12)
        self.show.grid(row=2, column=0, pady=3)
        self.rlt.grid(row=3, column=0, pady=3)
        self.manage.grid(row=4, column=0, pady=3)

        self.base.grid_propagate(0)
        self.show.grid_propagate(0)
        self.rlt.grid_propagate(0)
        self.manage.grid_propagate(0)

        # 初始化函数区
        self.showconfnum_func()
        self.showcondition_func()
        self.bind_func()

        self.root.mainloop()

    # 绑定各种数字的刷新
    def bind_func(self):
        self.bossnum_cb.bind('<<ComboboxSelected>>', self.showconfnum_func)
        self.confnum_cb.bind('<<ComboboxSelected>>', self.showconfnum_func)
        self.backexpertnum_cb.bind('<<ComboboxSelected>>', self.showconfnum_func)

        self.choicefield_cb.bind('<<ComboboxSelected>>', self.showcondition_func)
        self.choicecondition1_cb.bind('<<ComboboxSelected>>', self.showcondition_func)
        self.rlttree.bind('<Double-1>', self.expertcheck_func)

    def showcondition_func(self, *args):
        try:
            self.dfexpert = read_sql('select * from tbexpert', self.dbconn)
            self.dfcata = read_sql('select * from tbcata', self.dbconn)
            if self.choicefield_var.get() != '专业类别':
                self.condtion_lb.grid_forget()
                self.choicecondition1_cb.grid_forget()
                self.choicecondition_lb.delete(0, END)
                for name in self.dfexpert[self.choicefield_var.get()].unique():
                    self.choicecondition_lb.insert('end', name)
            else:
                self.condtion_lb.grid(row=1, column=2)
                self.choicecondition1_cb.grid(row=1, column=3, sticky=self.duiqi)
                self.choicecondition1_cb['values'] = list(self.dfcata['专业分类'].unique())
                self.choicecondition_lb.delete(0, END)
                for major in self.dfcata[self.dfcata['专业分类'] == self.choicecondition1_var.get()]['专业名称'].unique():
                    self.choicecondition_lb.insert('end', major)
        except:
            messagebox.showwarning(title='专家库', message='专家库异常，请检查专家库！')

    def addcondition_func(self):
        try:
            for cdn in self.choicecondition_lb.curselection():
                self.conditiontre_dict.setdefault(self.choicefield_var.get(), set())
                self.conditiontre_dict[self.choicefield_var.get()].add(self.choicecondition_lb.get(cdn))

            self.showconditiontree_func()
        except:
            messagebox.showinfo(title='条件选择有误', message='请选择正确条件！')

    def delcondition_func(self):
        # for x in self.tree.selection():
        #     self.tree.delete(x)
        self.conditiontre_dict = {}
        self.cleartree_func(self.tree)

        # 表格内容插入

    def showconditiontree_func(self):
        self.cleartree_func(self.tree)
        # print(self.conditiontre_dict)
        for key, values in self.conditiontre_dict.items():
            for value in values:
                self.tree.insert("", "end", values=(key, value))

    def showrlttree_func(self):
        self.cleartree_func(self.rlttree)
        for index, row in self.dfrltexpert.iterrows():
            self.rlttree.insert("", "end", values=(
                index, row['姓名'], row['性别'], row['工作单位'], row['工作部门'], row['区域'], row['专业资质等级'], row['是否参加'],
                row['抽取批次']))
        pass

    # 清空树
    def cleartree_func(self, tree):
        x = tree.get_children()
        for item in x:
            tree.delete(item)

    # 抽取数据的函数
    def choiceexpert_func(self):
        try:
            if int(self.expertchoicenum_var.get()) > 0:
                self.dfexpert = read_sql('select * from tbexpert', self.dbconn)
                # 不为空
                if self.dfexpert.shape[0] != 0:
                    self.dfconditionexpert = self.dfexpert.copy()
                    self.dfmajorexpert = self.dfexpert.set_index(['index']).stack().reset_index()
                    for key in self.conditiontre_dict.keys():
                        if key != '专业类别':
                            self.dfcondition = DataFrame(list(self.conditiontre_dict[key]), columns=[key])
                            self.dfconditionexpert = merge(self.dfcondition, self.dfconditionexpert, on=key,
                                                           how='inner')
                        else:
                            self.dfmajor = DataFrame(list(self.conditiontre_dict[key]), columns=[key])
                            self.dfmajor = merge(self.dfmajor, self.dfcata, left_on=key, right_on='专业名称', how='inner')
                            self.dfmajorexpert = merge(self.dfmajor, self.dfmajorexpert,
                                                       left_on='专业编号', right_on=0)

                    self.dfsampleexpert = merge(self.dfconditionexpert, self.dfmajorexpert, on='index', how='inner')
                    self.dfsampleexpert = DataFrame(self.dfsampleexpert['index'].unique(), columns=['index'])
                    # 去除重复的，已经抽取的就减掉。
                    _ = self.dfrltexpert.index.astype(int)
                    self.dfsampleexpert.set_index('index', inplace=True)
                    self.dfsampleexpert.drop(_, inplace=True)
                    self.dfsampleexpert.reset_index(inplace=True)

                    # dfsampleexpert是可以用来抽取的专家index清单，只有一列，sample就是在index，防止重名的情况发生
                    # dfrltexpert就是抽取出来的专家库，包含完整的列。
                    if self.dfsampleexpert.shape[0] > int(self.expertnum_var.get()):

                        # 折腾索引
                        _ = merge(self.dfsampleexpert.sample(int(self.expertchoicenum_var.get()), replace=False),
                                  self.dfexpert, on='index', how='inner')
                        _['index'] = _['index'].astype(str)
                        _.set_index('index', inplace=True)
                        _['是否参加'] = '未联系'
                        # _['是否参加'] = '参加'
                        self.choicetime_int += 1
                        _['抽取批次'] = self.choicetime_int
                        self.dfrltexpert = concat([self.dfrltexpert, _])
                        self.rlt_lb.config(
                            text='目前条件专家总数：' + str(self.dfsampleexpert.shape[0]) + '。第{}批次抽取{}个专家，'.format(
                                self.choicetime_int, _.shape[0]) + '已抽取专家人数：{}。'.format(str(self.dfrltexpert.shape[0])))
                        self.showrlttree_func()

                    else:
                        messagebox.showwarning(title='不够数量', message='目前条件专家总数：' + str(
                            self.dfsampleexpert.shape[0]) + '。\n专家数量不够，请重设条件！')
                else:
                    messagebox.showwarning(title='没有专家数据', message='没有专家数据，请导出专家！')
            else:
                messagebox.showinfo(title='抽取完成', message='如需重新抽取请重置结果！')
            # print(self.dfrltexpert)
            # 刷新数量
            self.showconfnum_func()
        except:
            messagebox.showerror(title='出错', message='程序出错，请与管理员联系！')

        # 重置结果

    def resetrlttree(self):
        self.dfrltexpert = DataFrame()
        self.choicetime_int = 0
        self.cleartree_func(self.rlttree)
        self.rlt_lb.config(text='尚未抽取专家！')
        self.showconfnum_func()

    def __del__(self):
        self.dbconn.close()

    # 导入专家库
    def importdata_func(self):
        filename = filedialog.askopenfilename()
        if filename != '':
            try:
                xlsfile = ExcelFile(filename)
                self.dfexpert = xlsfile.parse(sheet_name='专家名单', header=0, converters={x: str for x in range(15, 75)})
                # 全部转换成文本
                self.dfexpert = self.dfexpert.astype(str)
                self.dfexpert.replace('nan', nan, inplace=True)
                self.dfexpert.replace('None', nan, inplace=True)
                self.dfexpert.to_sql('tbexpert', self.dbconn, if_exists='replace')

                # 分类表
                self.dfcata = xlsfile.parse(sheet_name='专业分类', header=0, converters={0: str, 2: str})
                self.dfcata['编号'].fillna(method='ffill', inplace=True)
                self.dfcata.set_index('编号', inplace=True)
                self.dfcata['专业分类'].fillna(method='ffill', inplace=True)
                self.dfcata['专业编号'].fillna(method='ffill', inplace=True)
                self.dfcata['专业名称'].fillna(method='ffill', inplace=True)
                # 全部转换成文本
                self.dfcata = self.dfcata.astype(str)
                self.dfcata.replace('nan', nan, inplace=True)
                self.dfcata.replace('None', nan, inplace=True)
                self.dfcata.to_sql('tbcata', self.dbconn, if_exists='replace')

                messagebox.showinfo(title='导入成功', message='导入成功，可以按查看专家按钮查看。！')
            except:
                messagebox.showwarning(title='导入失败', message='导入失败，请重新导入！')

    # 查看历史
    def showhistory_func(self):
        historywindow(self)

    # 导出抽取结果数据
    def exportrlt_func(self):
        try:
            if int(self.expertchoicenum_var.get()) == 0 and self.dfrltexpert[self.dfrltexpert['是否参加'] == '未联系'].shape[
                0] == 0:
                # 此处添加导出文件的代码
                summary_doc = Document('评标专家抽取过程纪要函(模板).docx')
                summary_doc.paragraphs[1].text = summary_doc.paragraphs[1].text + self.projectname_var.get()
                summary_doc.paragraphs[2].text = summary_doc.paragraphs[2].text + self.projectid_var.get()
                summary_doc.paragraphs[3].text = summary_doc.paragraphs[3].text + self.meetdate_var.get()
                summary_doc.paragraphs[4].text = summary_doc.paragraphs[4].text + self.choiceplace_var.get()
                _ = summary_doc.paragraphs[5].text.split('|')
                summary_doc.paragraphs[5].text = _[0] + self.choicedate_var.get() + _[1] + self.choiceplace_var.get() + \
                                                 _[2]
                _ = summary_doc.paragraphs[6].text.split('|')
                summary_doc.paragraphs[6].text = _[0] + self.confnum_var.get() + _[1] + self.bossnum_var.get() + _[
                    2] + str(int(self.confnum_var.get()) - int(self.bossnum_var.get())) + _[3]
                summary_doc.paragraphs[7].text = summary_doc.paragraphs[7].text
                summary_doc.paragraphs[8].text = summary_doc.paragraphs[8].text

                i = 1
                for index, row in self.dfrltexpert.iterrows():
                    summary_doc.tables[0].add_row()
                    summary_doc.tables[0].rows[i].cells[0].text = str(i)
                    summary_doc.tables[0].rows[i].cells[1].text = row['姓名']
                    summary_doc.tables[0].rows[i].cells[2].text = '{}/{}'.format(row['工作单位'], row['工作部门'])
                    summary_doc.tables[0].rows[i].cells[3].text = row['联系电话(133)']
                    summary_doc.tables[0].rows[i].cells[4].text = str(row['抽取批次'])
                    summary_doc.tables[0].rows[i].cells[5].text = row['是否参加']
                    i = i + 1

                dfexpertfinal = self.dfrltexpert[self.dfrltexpert['是否参加'] == '参加']

                # 打开专家费签收表
                wb = xlrd.open_workbook('专家费签收表(模板).xls', formatting_info=True)
                wsread = wb.sheet_by_index(0)
                wbsave = copy(wb)
                ws = wbsave.get_sheet(0)
                # 写入的样式
                wsstyle = xlwt.easyxf('font: bold on; align: wrap on, vert centre, horiz center')
                wsstyle.font.height = 220
                wsstyle.borders.top = 1
                wsstyle.borders.bottom = 1
                wsstyle.borders.left = 1
                wsstyle.borders.right = 1

                j = 1
                for index, row in dfexpertfinal.iterrows():
                    summary_doc.tables[1].add_row()
                    summary_doc.tables[1].rows[j].cells[0].text = str(j)
                    summary_doc.tables[1].rows[j].cells[1].text = row['姓名']
                    summary_doc.tables[1].rows[j].cells[2].text = '{}/{}'.format(row['工作单位'], row['工作部门'])
                    summary_doc.tables[1].rows[j].cells[3].text = '{}/{}'.format(row['联系电话(133)'], row['E-MAIL'])
                    summary_doc.tables[1].rows[j].cells[4].text = '技术/经济'
                    summary_doc.tables[1].rows[j].cells[5].text = '抽取'

                    # 正式写入给钱的信息
                    ws.write(j + 5, 1, row['姓名'], wsstyle)
                    ws.write(j + 5, 3, '{}/{}'.format(row['工作单位'], row['工作部门']), wsstyle)
                    ws.write(j + 5, 4, row['身份证号'], wsstyle)
                    j = j + 1

                # 如果有业主代表
                for bossindex in range(int(self.bossnum_var.get())):
                    summary_doc.tables[1].add_row()
                    summary_doc.tables[1].rows[j].cells[0].text = str(j)
                    summary_doc.tables[1].rows[j].cells[1].text = self.dfboss.loc[bossindex, 'bossname']
                    summary_doc.tables[1].rows[j].cells[2].text = self.dfboss.loc[bossindex, 'bossworkpart']
                    summary_doc.tables[1].rows[j].cells[3].text = '{}/{}'.format(
                        self.dfboss.loc[bossindex, 'bossmobilno'], self.dfboss.loc[bossindex, 'bossemail'])
                    summary_doc.tables[1].rows[j].cells[4].text = '业主代表'
                    summary_doc.tables[1].rows[j].cells[5].text = '推荐'

                    # 正式写入给钱的信息
                    ws.write(j + 5, 1, self.dfboss.loc[bossindex, 'bossname'], wsstyle)
                    ws.write(j + 5, 3, self.dfboss.loc[bossindex, 'bossworkpart'], wsstyle)
                    ws.write(j + 5, 4, self.dfboss.loc[bossindex, 'bossnaid'], wsstyle)
                    j = j + 1

                dfexpertfinal = dfexpertfinal.append(pd.DataFrame(
                    {'姓名': self.dfboss.iloc[:int(self.bossnum_var.get()), :]['bossname'], '工作部门':
                        self.dfboss.iloc[:int(self.bossnum_var.get()), :]['bossworkpart'], '身份证号':
                         self.dfboss.iloc[:int(self.bossnum_var.get()), :]['bossnaid'], '联系电话(办公)':
                         self.dfboss.iloc[:int(self.bossnum_var.get()), :]['bosstelno'], '联系电话(133)':
                         self.dfboss.iloc[:int(self.bossnum_var.get()), :]['bossmobilno'], 'E-MAIL':
                         self.dfboss.iloc[:int(self.bossnum_var.get()), :]['bossemail'], '工作单位': '业主单位','性别':'未知'}),
                    ignore_index=True)

                # 保存word
                summary_doc.paragraphs[13].text = self.choicedate_var.get()
                summary_doc.save('评标专家抽取过程纪要函' + dt.today().strftime('%Y%m%d') + '.docx')
                # excel写入
                txts = wsread.cell(2, 0).value.split('|')
                txt = '{}'.join(txts)
                ws.write(2, 0, txt.format(self.projectname_var.get(), self.projectid_var.get()),
                         xlwt.easyxf('font: bold on; align: wrap on, vert centre, horiz left'))
                ws.write(3, 8, '评审时间：{}'.format(self.meetdate_var.get()),
                         xlwt.easyxf('font: bold on; align: vert centre, horiz right'))
                wbsave.save('专家费签收表' + dt.today().strftime('%Y%m%d') + '.xls')

                # 写入数据库
                dftosql = dfexpertfinal.copy()
                prjkey = ''
                if self.projectname_var.get() == '':
                    prjkey = '空项目'
                else:
                    prjkey = self.projectname_var.get()
                dftosql['项目名称'] = prjkey
                dftosql['项目编号'] = self.projectid_var.get()
                dftosql['招标人名称'] = self.bossname_var.get()
                dftosql.reset_index(inplace=True)
                dftosql.set_index('项目名称', inplace=True)
                try:
                    # 如果有重复，加入前先删除,一个项目名称只能有一次抽取记录被保存。
                    if self.dbcur.execute('select count(*) from tbhistory where 项目名称=?', (prjkey,)).fetchall()[0][
                        0] > 0:
                        self.dbcur.execute('delete from tbhistory where 项目名称=?', (prjkey,))
                        self.dbconn.commit()
                except:
                    pass
                dftosql.to_sql('tbhistory', self.dbconn, if_exists='append')
                messagebox.showinfo(title='专家抽取纪要生成', message='已经生成专家抽取纪要，请在程序安装目录查看！')
            else:
                messagebox.showinfo(title='专家数量不符合要求', message='请确定参加评标会议的专家数量符合要求！')

        except BaseException as e:
            print(e)
            messagebox.showwarning(title='警告', message='导出错误，请联系管理员！')

    # 清空专家库
    def cleardata_func(self):
        self.dfnull = DataFrame()
        self.dfcata = self.dfnull
        self.dfexpert = self.dfnull
        self.dfnull.to_sql('tbexpert', self.dbconn, if_exists='replace')
        self.dfnull.to_sql('tbcata', self.dbconn, if_exists='replace')
        # self.dfnull.to_sql('tbhistory', self.dbconn, if_exists='replace')
        messagebox.showinfo(title='清空专家库', message='专家库已经清空！')

    def checknum_func(self):
        try:
            int(self.confnum_var.get())
        except:
            messagebox.showwarning(title='注意评标委员为人数', message='请填入3以上单数')
            return False
        if self.confnum_var.get() == '':
            messagebox.showwarning(title='注意评标委员为人数', message='不能为空')
            return False
        if int(self.confnum_var.get()) <= 3 or int(self.confnum_var.get()) % 2 == 0:
            messagebox.showwarning(title='注意评标委员为人数', message='不符合3人以上单数的要求')
            return False
        if self.bossnum_var.get() == '':
            messagebox.showwarning(title='注意业主代表为人数', message='请选择业主代表人数')
            return False
        if self.backexpertnum_var.get() == '':
            messagebox.showwarning(title='注意递补专家为人数', message='请选择递补专家人数')
            return False
        return True

    def checktext(self):
        pass

    def showconfnum_func(self, *args):
        if self.checknum_func():
            self.expertnum_var.set(int(self.confnum_var.get()) - int(self.bossnum_var.get()))
            try:
                joinexpert = \
                    self.dfrltexpert[(self.dfrltexpert['是否参加'] == '参加') | (self.dfrltexpert['是否参加'] == '未联系')].shape[0]
            except:
                joinexpert = 0
            self.expertchoicenum_var.set(int(self.confnum_var.get()) + int(self.backexpertnum_var.get()) - int(
                self.bossnum_var.get()) - joinexpert)

    # 查看专家
    def showexpert_func(self):
        try:
            expertwindow(self)
        except:
            messagebox.showerror(title='专家库错误', message='专家库错误，请联系管理员！')

    # 确定专家
    def expertcheck_func(self, *args):
        try:
            expertcheck(self)
        except:
            messagebox.showwarning(title='请抽取专家', message='请抽取专家')

    # 输入业主代表的信息
    def bossinfo_func(self):
        bosswindow(self)


class expertwindow():
    def __init__(self, mother):
        self.mother = mother
        self.top = Toplevel(self.mother.root, width=900, height=500)
        self.top.title('专家库')
        self.top.iconbitmap('logo.ico')
        # self.top.attributes('-topmost', 1)

        # 基本信息
        self.expertnum_var = StringVar()
        Label(self.top, text='专家数量').grid(row=0, column=0)
        self.expertnum_et = Entry(self.top, state='readonly', textvariable=self.expertnum_var).grid(row=0, column=1)
        Button(self.top, text='导出数据', command=self.exportexpert_func, width=30).grid(row=0, column=2)
        Button(self.top, text='退出', command=self.top.destroy, width=30).grid(row=0, column=3)

        # 初始化树
        self.experttree = ttk.Treeview(self.top, show="headings", height=20,
                                       columns=("a", "b", "c", "d", "e", "f", 'g'))
        self.expertvbar = ttk.Scrollbar(self.top, orient=VERTICAL, command=self.experttree.yview)
        # 定义树形结构与滚动条
        self.experttree.configure(yscrollcommand=self.expertvbar.set)
        # 表格的标题
        self.experttree.column("a", width=80, anchor="center")
        self.experttree.column("b", width=50, anchor="center")
        self.experttree.column("c", width=190, anchor="center")
        self.experttree.column("d", width=190, anchor="center")
        self.experttree.column("e", width=50, anchor="center")
        self.experttree.column("f", width=190, anchor="center")
        self.experttree.column("g", width=180, anchor="center")
        self.experttree.heading("a", text="姓名")
        self.experttree.heading("b", text="性别")
        self.experttree.heading("c", text='工作单位')
        self.experttree.heading("d", text="工作部门")
        self.experttree.heading("e", text="区域")
        self.experttree.heading("f", text="资质等级")
        self.experttree.heading("g", text="身份证号")
        # 调用方法获取表格内容插入
        self.experttree.grid(row=1, column=0, sticky='NEW', columnspan=4)
        self.expertvbar.grid(row=1, column=4, sticky='NS', columnspan=4)

        # 显示专家树
        self.showexperttree_func()

    def __del__(self):
        pass

    # 显示专家树
    def showexperttree_func(self):

        try:
            self.dfexpert = read_sql('select * from tbexpert', self.mother.dbconn)
            self.dfcata = read_sql('select * from tbcata', self.mother.dbconn)
            self.expertnum_var.set(self.dfexpert['姓名'].count())
            for index, row in self.dfexpert.iterrows():
                self.experttree.insert("", "end", values=(
                    row['姓名'], row['性别'], row['工作单位'], row['工作部门'], row['区域'], row['专业资质等级'], row['身份证号']))
        except:
            self.experttree.insert("", "end", values=('没有数据', '', '', '', '', ''))

    # 导出专家库
    def exportexpert_func(self):
        try:
            writer = ExcelWriter('导出专家库' + dt.today().strftime('%Y%m%d') + '.xls')
            self.dfexpert.iloc[:, 1:].to_excel(writer, sheet_name='专家名单', index=False)
            self.dfcata.to_excel(writer, sheet_name='专业分类', index=False)
            writer.save()
            messagebox.showinfo(title='导出成功', message='导出成功，请在程序同一个文件夹查看。')

        except:
            messagebox.showwarning(title='导出失败', message='导出失败，请重新导入！')


# 历史窗口
class historywindow():
    def __init__(self, mother):
        self.mother = mother
        self.top = Toplevel(self.mother.root, width=900, height=500)
        self.top.title('抽取历史')
        self.top.iconbitmap('logo.ico')
        # self.top.attributes('-topmost', 1)

        # 基本信息
        self.historynum_var = StringVar()
        Label(self.top, text='抽取项目数').grid(row=0, column=0)
        self.historynum_var_et = Entry(self.top, state='readonly', textvariable=self.historynum_var).grid(row=0,
                                                                                                          column=1)
        Button(self.top, text='导出数据', command=self.historyexpert_func, width=25).grid(row=0, column=2)
        Button(self.top, text='清空数据', command=self.cleardata_func, width=25).grid(row=0, column=3)
        Button(self.top, text='退出', command=self.top.destroy, width=25).grid(row=0, column=4)

        # 初始化树
        self.historytree = ttk.Treeview(self.top, show="headings", height=20,
                                        columns=("a", "b", "c", "d", "e", "f", "g"))
        self.historyvbar = ttk.Scrollbar(self.top, orient=VERTICAL, command=self.historytree.yview)
        # 定义树形结构与滚动条
        self.historytree.configure(yscrollcommand=self.historyvbar.set)
        # 表格的标题
        self.historytree.column("a", width=80, anchor="center")
        self.historytree.column("b", width=50, anchor="center")
        self.historytree.column("c", width=160, anchor="center")
        self.historytree.column("d", width=140, anchor="center")
        self.historytree.column("e", width=180, anchor="center")
        self.historytree.column("f", width=120, anchor="center")
        self.historytree.column("g", width=120, anchor="center")
        self.historytree.heading("a", text="姓名")
        self.historytree.heading("b", text="性别")
        self.historytree.heading("c", text='工作单位')
        self.historytree.heading("d", text="工作部门")
        self.historytree.heading("e", text="身份证号")
        self.historytree.heading("f", text="项目名称")
        self.historytree.heading("g", text="招标人名称")
        # 调用方法获取表格内容插入
        self.historytree.grid(row=1, column=0, sticky='NEW', columnspan=5)
        self.historyvbar.grid(row=1, column=5, sticky='NS', columnspan=4)

        # 显示历史树
        self.showhistorytree_func()

    def __del__(self):
        pass

    # 显示历史树
    def showhistorytree_func(self):
        try:
            self.dfhistory = read_sql('select 项目名称,项目编号,招标人名称,代理序号,姓名,性别,工作单位,工作部门,身份证号 from tbhistory',
                                      self.mother.dbconn)
            self.historynum_var.set(self.dfhistory['项目名称'].unique().shape[0])
            for index, row in self.dfhistory.iterrows():
                self.historytree.insert("", "end", values=(
                    row['姓名'], row['性别'], row['工作单位'], row['工作部门'], row['身份证号'], row['项目名称'], row['招标人名称']))
        except:
            self.historytree.insert("", "end", values=('没有数据', '', '', '', '', ''))

    # 导出历史数据库
    def historyexpert_func(self):
        try:
            writer = ExcelWriter('导出历史数据' + dt.today().strftime('%Y%m%d') + '.xls')
            self.dfhistory.to_excel(writer, sheet_name='抽取历史数据', index=False)
            writer.save()
            messagebox.showinfo(title='导出成功', message='导出成功，请在程序同一个文件夹查看。')

        except:
            messagebox.showwarning(title='导出失败', message='导出失败，请重新导入！')

    # 清空数据
    def cleardata_func(self):
        try:
            self.mother.dbconn.execute('drop table tbhistory')
            self.mother.dbconn.commit()
            messagebox.showinfo(title='清空历史数据', message='历史数据已经清空！')
            self.mother.cleartree_func(self.historytree)
        except:
            pass

    def cleartree_func(self, tree):
        x = tree.get_children()
        for item in x:
            tree.delete(item)


class expertcheck():
    def __init__(self, mother):
        self.mother = mother
        self.top = Toplevel(self.mother.root, width=600, height=300)
        self.top.title('专家确定')
        self.top.geometry('600x200')
        self.top.iconbitmap('logo.ico')

        self.stdwidth = 14
        self.etwidth = 28
        self.expertinfo = Frame(self.top, width=600, height=200)
        self.expertinfo.grid(row=0, column=0)
        self.expertindex = str(self.mother.rlttree.item(mother.rlttree.focus())['values'][0])
        self.dfexpert = self.mother.dfrltexpert.loc[self.expertindex, :]
        Label(self.expertinfo, text='姓名：', width=self.stdwidth).grid(row=0, column=0, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['姓名'], width=self.etwidth).grid(row=0, column=1, sticky=W)
        Label(self.expertinfo, text='性别：', width=self.stdwidth).grid(row=0, column=2, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['性别'], width=self.etwidth).grid(row=0, column=3, sticky=W)
        Label(self.expertinfo, text='现任职务：', width=self.stdwidth).grid(row=1, column=0, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['现任职务'], width=self.etwidth).grid(row=1, column=1, sticky=W)
        Label(self.expertinfo, text='工作单位：', width=self.stdwidth).grid(row=1, column=2, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['工作单位'], width=self.etwidth).grid(row=1, column=3, sticky=W)
        Label(self.expertinfo, text='工作部门：', width=self.stdwidth).grid(row=2, column=0, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['工作部门'], width=self.etwidth).grid(row=2, column=1, sticky=W)
        Label(self.expertinfo, text='区域：', width=self.stdwidth).grid(row=2, column=2, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['区域'], width=self.etwidth).grid(row=2, column=3, sticky=W)
        Label(self.expertinfo, text='联系电话(办公)：', width=self.stdwidth).grid(row=3, column=0, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['联系电话(办公)'], width=self.etwidth).grid(row=3, column=1, sticky=W)
        Label(self.expertinfo, text='联系电话(手机)：', width=self.stdwidth).grid(row=3, column=2, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['联系电话(133)'], width=self.etwidth).grid(row=3, column=3, sticky=W)
        Label(self.expertinfo, text='E-MAIL：', width=self.stdwidth).grid(row=4, column=0, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['E-MAIL'], width=self.etwidth).grid(row=4, column=1, sticky=W)
        Label(self.expertinfo, text='专业资质名称：', width=self.stdwidth).grid(row=4, column=2, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['专业资质名称'], width=self.etwidth).grid(row=4, column=3, sticky=W)
        Label(self.expertinfo, text='专业资质等级：', width=self.stdwidth).grid(row=5, column=0, sticky=W)
        Label(self.expertinfo, text=self.dfexpert['专业资质等级'], width=self.etwidth).grid(row=5, column=1, sticky=W)
        # 选择
        Label(self.expertinfo, text='是否参加：', width=self.stdwidth).grid(row=5, column=2, sticky=W)
        self.isjoin_var = StringVar()
        self.isjoin_cb = ttk.Combobox(self.expertinfo, height=1, state='readonly', textvariable=self.isjoin_var,
                                      width=int(self.etwidth * 0.6))
        self.isjoin_cb['values'] = ('参加', '不参加', '无法联系', '回避')
        # self.isjoin_cb.current(0)
        self.isjoin_var.set(self.dfexpert['是否参加'])
        self.isjoin_cb.grid(row=5, column=3, sticky=W)

        # 备注
        # Label(self.expertinfo, text='备注：').grid(row=6, column=0, sticky=W)
        # Text(self.expertinfo,height=4).grid(row=7,column=0,columnspan=4)

        Button(self.expertinfo, text='确定', width=20, command=self.expertconfirm_func).grid(row=6, column=0, pady=20,
                                                                                           columnspan=2)
        Button(self.expertinfo, text='取消', width=20, command=lambda: self.top.destroy()).grid(row=6, column=2, pady=20,
                                                                                              columnspan=2)

    def expertconfirm_func(self):
        self.mother.dfrltexpert.loc[self.expertindex, '是否参加'] = self.isjoin_var.get()
        self.mother.showrlttree_func()
        self.mother.showconfnum_func()
        self.top.destroy()


# 业主代表信息
class bosswindow():
    def __init__(self, mother):
        self.mother = mother
        self.top = Toplevel(self.mother.root, width=900, height=300)
        self.top.title('推荐专家信息')
        self.top.iconbitmap('logo.ico')
        # self.top.attributes('-topmost', 1)

        # 基本信息
        self.bossnum_var = StringVar()
        Label(self.top, text='推荐专家数量').grid(row=0, column=0)
        self.historynum_var_et = Entry(self.top, state='readonly', textvariable=self.bossnum_var).grid(row=0,
                                                                                                       column=1)
        Button(self.top, text='确定', command=self.confirm_func, width=25).grid(row=0, column=2)

        # 初始化树
        self.bosstree = ttk.Treeview(self.top, show="headings", height=6,
                                     columns=("a", "b", "c", "d", "e", "f", "g"))
        self.bossvbar = ttk.Scrollbar(self.top, orient=VERTICAL, command=self.bosstree.yview)
        # 定义树形结构与滚动条
        self.bosstree.configure(yscrollcommand=self.bossvbar.set)
        # 表格的标题
        self.bosstree.column("a", width=50, anchor="center")
        self.bosstree.column("b", width=50, anchor="center")
        self.bosstree.column("c", width=160, anchor="center")
        self.bosstree.column("d", width=160, anchor="center")
        self.bosstree.column("e", width=160, anchor="center")
        self.bosstree.column("f", width=160, anchor="center")
        self.bosstree.column("g", width=160, anchor="center")

        self.bosstree.heading("a", text="序号")
        self.bosstree.heading("b", text="姓名")
        self.bosstree.heading("c", text="工作单位")
        self.bosstree.heading("d", text='工作电话')
        self.bosstree.heading("e", text="手机")
        self.bosstree.heading("f", text="EMAIL")
        self.bosstree.heading("g", text="身份证号")

        # 调用方法获取表格内容插入
        self.bosstree.grid(row=1, column=0, sticky='NEW', columnspan=5)
        self.bossvbar.grid(row=1, column=5, sticky='NS', columnspan=4)

        # 显示历史树
        self.showhbosstree_func()
        self.bosstree.bind('<Double-1>', self.bosscheck_func)

    def __del__(self):
        pass

    # 显示历史树
    def showhbosstree_func(self):
        try:
            self.cleartree_func(self.bosstree)
            self.bossnum_var.set(self.mother.bossnum_var.get())
            for index, row in self.mother.dfboss.iterrows():
                if index >= int(self.mother.bossnum_var.get()):
                    break
                self.bosstree.insert("", "end", values=(index,
                                                        row['bossname'], row['bossworkpart'], row['bosstelno'],
                                                        row['bossmobilno'], row['bossemail'],
                                                        row['bossnaid']))
        except:
            self.bosstree.insert("", "end", values=('没有数据', '', '', '', '', ''))

    def confirm_func(self):
        try:
            self.mother.bossname_lb.config(
                text='/'.join(list(self.mother.dfboss.iloc[:int(self.mother.bossnum_var.get()), :]['bossname'])))
        except BaseException as e:
            print(e)
            self.mother.bossname_lb.config(text='无推荐专家')
        self.top.destroy()

    # 确定业主代表
    def bosscheck_func(self, *args):
        try:
            bossinfo(self)
        except BaseException as e:
            print(e)
            messagebox.showwarning(title='出错', message='请联系管理员')

    def cleartree_func(self, tree):
        x = tree.get_children()
        for item in x:
            tree.delete(item)


class bossinfo():
    def __init__(self, mother):
        self.mother = mother
        self.top = Toplevel(self.mother.top, width=500, height=300)
        self.top.title('业主代表信息')
        self.top.geometry('400x150')
        self.top.iconbitmap('logo.ico')
        self.bossinfo = Frame(self.top, width=500, height=200)
        self.bossinfo.grid(row=0, column=0, pady=20)
        self.bosssel = self.mother.bosstree.item(mother.bosstree.focus())['values']

        # 业主代表姓名
        Label(self.bossinfo, text='姓名：', width=8).grid(row=0, column=0, sticky=W)
        self.bossname_var = StringVar()
        Entry(self.bossinfo, width=18, textvariable=self.bossname_var).grid(row=0, column=1, sticky=W)

        # 工作部门
        Label(self.bossinfo, text='工作部门：', width=8).grid(row=0, column=2, sticky=W)
        self.workpart_var = StringVar()
        Entry(self.bossinfo, width=18, textvariable=self.workpart_var).grid(row=0, column=3, sticky=W)

        # 工作电话
        Label(self.bossinfo, text='工作电话：', width=8).grid(row=1, column=0, sticky=W)
        self.telno_var = StringVar()
        Entry(self.bossinfo, width=18, textvariable=self.telno_var).grid(row=1, column=1, sticky=W)

        # 工作手机
        Label(self.bossinfo, text='工作手机：', width=8).grid(row=1, column=2, sticky=W)
        self.mobilno_var = StringVar()
        Entry(self.bossinfo, width=18, textvariable=self.mobilno_var).grid(row=1, column=3, sticky=W)

        # 工作邮箱
        Label(self.bossinfo, text='工作邮箱：', width=8).grid(row=2, column=0, sticky=W)
        self.email_var = StringVar()
        Entry(self.bossinfo, width=18, textvariable=self.email_var).grid(row=2, column=1, sticky=W)

        # 身份号
        Label(self.bossinfo, text='身份证号：', width=8).grid(row=2, column=2, sticky=W)
        self.naid_var = StringVar()
        Entry(self.bossinfo, width=18, textvariable=self.naid_var).grid(row=2, column=3, sticky=W)

        # confirn button
        Button(self.bossinfo, text='确定', width=18, command=self.bossconfirm_func).grid(row=3, column=0, columnspan=2,
                                                                                       pady=20, sticky=S)
        Button(self.bossinfo, text='取消', width=18, command=lambda: self.top.destroy()).grid(row=3, column=2,
                                                                                            columnspan=2, pady=20,
                                                                                            sticky=S)

        # 初始化数据
        self.bossname_var.set(self.bosssel[1])
        self.workpart_var.set(self.bosssel[2])
        self.telno_var.set(self.bosssel[3])
        self.mobilno_var.set(self.bosssel[4])
        self.email_var.set(self.bosssel[5])
        self.naid_var.set(self.bosssel[6])

    def bossconfirm_func(self):
        self.mother.mother.dfboss.loc[self.bosssel[0], 'bossname'] = self.bossname_var.get()
        self.mother.mother.dfboss.loc[self.bosssel[0], 'bossworkpart'] = self.workpart_var.get()
        self.mother.mother.dfboss.loc[self.bosssel[0], 'bosstelno'] = self.telno_var.get()
        self.mother.mother.dfboss.loc[self.bosssel[0], 'bossmobilno'] = self.mobilno_var.get()
        self.mother.mother.dfboss.loc[self.bosssel[0], 'bossemail'] = self.email_var.get()
        self.mother.mother.dfboss.loc[self.bosssel[0], 'bossnaid'] = self.naid_var.get()

        self.mother.showhbosstree_func()
        self.top.destroy()


if __name__ == '__main__':
    expertChoice()
